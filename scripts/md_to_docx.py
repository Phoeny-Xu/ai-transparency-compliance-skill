# -*- coding: utf-8 -*-
"""通用 Markdown -> Word(docx) 转换器，针对本报告类文档优化。
增强版：支持真实 Word 脚注（word/footnotes.xml）、表格表头灰底、超链接、
批注章节转 Word 批注气泡。脚注语法：正文 [^label]，定义行 [^label]: 文本。
"""
import re
import sys
import os
import shutil
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 中文字体设置 ----------
CN_FONT = "宋体"
CN_FONT_BOLD = "黑体"
HEADING_FONT = "宋体"

# 脚注全局映射（convert 填充，add_footnote_ref / _inject_footnotes 读取）
FOOTNOTE_DEFS = {}
FOOTNOTE_LABEL_TO_ID = {}


def set_cn_font(run, font_name=CN_FONT, size=None, bold=False):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rfonts.set(qn('w:ascii'), CN_FONT)
    rfonts.set(qn('w:hAnsi'), CN_FONT)
    rPr.append(rfonts)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ---------- 内联解析（**加粗** / [文本](url) / 裸url / <font color> / <span style> / [^脚注]） ----------
def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    if len(hex_str) == 3:
        hex_str = ''.join(c * 2 for c in hex_str)
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def add_footnote_ref(paragraph, label):
    fid = FOOTNOTE_LABEL_TO_ID.get(label)
    if fid is None:
        run = paragraph.add_run('[^%s]' % label)
        set_cn_font(run, CN_FONT)
        return
    run = paragraph.add_run()
    rPr = run._element.get_or_add_rPr()
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    fr = OxmlElement('w:footnoteReference')
    fr.set(qn('w:id'), str(fid))
    run._element.append(fr)


def _add_plain_segment(paragraph, text):
    """处理普通文本段中的 **bold** / 裸url / [^脚注]"""
    parts = re.split(r'(\*\*[^*]+\*\*|https?://\S+|\[\^([^\]]+)\])', text)
    for part in parts:
        if not part:
            continue
        fm = re.match(r'\[\^([^\]]+)\]$', part)
        if fm:
            add_footnote_ref(paragraph, fm.group(1))
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_cn_font(run, CN_FONT, bold=True)
        elif re.match(r'^https?://\S+$', part):
            add_hyperlink(paragraph, part, part)
        else:
            run = paragraph.add_run(part)
            set_cn_font(run, CN_FONT)


def parse_inline(paragraph, text):
    """内联解析：支持 [text](url)、**bold**、裸url、<font color>、<span style> 与 [^脚注]"""
    pattern = re.compile(
        r'<font\s+color="(#?[0-9A-Fa-f]+)"\s*>(.*?)</font>'
        r'|<span\s+style="([^"]*)"\s*>(.*?)</span>'
        r'|\[([^\]]+)\]\((https?://[^)]+)\)'
        r'|\*\*(.+?)\*\*'
        r'|\[\^([^\]]+)\]'
        r'|(https?://\S+)',
        re.DOTALL,
    )
    last = 0
    n = len(text)
    for m in pattern.finditer(text):
        if m.start() > last:
            _add_plain_segment(paragraph, text[last:m.start()])
        last = m.end()
        if m.group(1) is not None:  # <font color>
            run = paragraph.add_run(m.group(2))
            run.font.color.rgb = hex_to_rgb(m.group(1))
            set_cn_font(run, CN_FONT)
        elif m.group(3) is not None:  # <span style>
            style = m.group(3)
            run = paragraph.add_run(m.group(4))
            cm = re.search(r'color:\s*(red|#?[0-9A-fa-f]+)', style)
            if cm:
                val = cm.group(1)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) if val == 'red' else hex_to_rgb(val)
            is_bold = 'bold' in style
            set_cn_font(run, CN_FONT, bold=is_bold)
        elif m.group(5) is not None:  # [text](url)
            add_hyperlink(paragraph, m.group(6), m.group(5))
        elif m.group(7) is not None:  # **bold**
            run = paragraph.add_run(m.group(7))
            set_cn_font(run, CN_FONT, bold=True)
        elif m.group(8) is not None:  # [^footnote]
            add_footnote_ref(paragraph, m.group(8))
        elif m.group(9) is not None:  # 裸url
            add_hyperlink(paragraph, m.group(9), m.group(9))
    if last < n:
        _add_plain_segment(paragraph, text[last:])


def is_table_separator(line):
    s = line.strip()
    if not s.startswith('|'):
        return False
    cleaned = s.replace('|', '').replace(' ', '')
    if cleaned == '':
        return False
    return set(cleaned) <= set('-:')


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells


def _is_serial(text):
    """判断单元格文本是否像序号（纯数字 / CP-X / 带圈数字 / 数字+标点）"""
    t = text.strip()
    return bool(re.match(r'^(\d+|CP-\d+|[①②③④⑤⑥⑦⑧⑨⑩]|\d+[\.、])$', t))


def style_table(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头：加粗 + 水平居中 + 垂直居中 + 浅灰底
    hdr = table.rows[0]
    for cell in hdr.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                set_cn_font(run, CN_FONT, size=9, bold=True)
        # 浅灰底
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E2F3')
        tcPr.append(shd)
    # 数据单元格：垂直居中；首列为序号时水平居中，其余水平左对齐
    for r_i, row in enumerate(table.rows[1:], start=1):
        for c_i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            is_serial = (c_i == 0 and _is_serial(cell.text))
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_serial else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_cn_font(run, CN_FONT, size=9, bold=run.font.bold)


def _inject_footnotes(docx_path):
    """将 FOOTNOTE_DEFS / FOOTNOTE_LABEL_TO_ID 注入为真实 Word 脚注。"""
    global FOOTNOTE_DEFS, FOOTNOTE_LABEL_TO_ID
    if not FOOTNOTE_DEFS:
        return 0
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    parts = []
    parts.append('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>')
    parts.append('<w:footnotes xmlns:w="%s" xmlns:r="%s">' % (ns_w, ns_r))
    parts.append('<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>')
    parts.append('<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>')
    for label, fid in sorted(FOOTNOTE_LABEL_TO_ID.items(), key=lambda kv: kv[1]):
        text = FOOTNOTE_DEFS.get(label, '')
        escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        parts.append('<w:footnote w:type="normal" w:id="%d">' % fid)
        parts.append('<w:p><w:pPr><w:pStyle w:val="FootnoteText"/><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>')
        parts.append('<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>')
        parts.append('<w:r><w:rPr><w:rFonts w:eastAsia="%s" w:ascii="%s" w:hAnsi="%s"/><w:sz w:val="18"/></w:rPr><w:t xml:space="preserve"> %s</w:t></w:r>' % (CN_FONT, CN_FONT, CN_FONT, escaped))
        parts.append('</w:p></w:footnote>')
    parts.append('</w:footnotes>')
    footnotes_xml = ''.join(parts).encode('utf-8')

    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    data['word/footnotes.xml'] = footnotes_xml
    # [Content_Types].xml
    ct = data['[Content_Types].xml'].decode('utf-8')
    if 'footnotes.xml' not in ct:
        ct = ct.replace('</Types>', '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
        data['[Content_Types].xml'] = ct.encode('utf-8')
    # document.xml.rels
    rels_name = 'word/_rels/document.xml.rels'
    if rels_name in data:
        rels = data[rels_name].decode('utf-8')
        if 'footnotes' not in rels:
            nums = [int(x) for x in re.findall(r'rId(\d+)', rels)]
            new_id = 'rId%d' % (max(nums) + 1) if nums else 'rId1'
            rels = rels.replace('</Relationships>',
                '<Relationship Id="%s" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>' % new_id)
            data[rels_name] = rels.encode('utf-8')
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for n, b in data.items():
            zout.writestr(n, b)
    shutil.move(tmp, docx_path)
    return len(FOOTNOTE_LABEL_TO_ID)


def convert(md_path, docx_path):
    global FOOTNOTE_DEFS, FOOTNOTE_LABEL_TO_ID
    FOOTNOTE_DEFS = {}
    FOOTNOTE_LABEL_TO_ID = {}
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    # 预扫脚注定义行 [^label]: text，并分配 id（从 1 起）
    fn_def_re = re.compile(r'^\[\^([^\]]+)\]:\s*(.*)$')
    for line in lines:
        m = fn_def_re.match(line.strip())
        if m:
            FOOTNOTE_DEFS[m.group(1)] = m.group(2)
    for idx, label in enumerate(FOOTNOTE_DEFS.keys(), start=1):
        FOOTNOTE_LABEL_TO_ID[label] = idx

    doc = Document()
    # 全局默认字体
    style = doc.styles['Normal']
    style.font.name = CN_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CN_FONT)
    rfonts.set(qn('w:ascii'), CN_FONT)
    rfonts.set(qn('w:hAnsi'), CN_FONT)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if stripped == '':
            i += 1
            continue

        # 脚注定义行：跳过（已收集）
        if fn_def_re.match(stripped):
            i += 1
            continue

        # 分隔线
        if stripped == '---' or stripped == '***':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            h = doc.add_heading(level=level)
            run = h.add_run(title)
            set_cn_font(run, HEADING_FONT, size=[18, 15, 13, 11.5][min(level-1, 3)], bold=True)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if level <= 2:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 引用块（可能多行）
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            # 合并连续引用
            j = i + 1
            while j < n and lines[j].strip().startswith('>'):
                quote_text += '\n' + lines[j].strip().lstrip('>').strip()
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            # 引用左侧竖线
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '4472C4')
            pBdr.append(left)
            pPr.append(pBdr)
            for k, seg in enumerate(quote_text.split('\n')):
                if k > 0:
                    p.add_run().add_break()
                parse_inline(p, seg)
            for run in p.runs:
                run.font.italic = True
                run.font.size = Pt(9.5)
                set_cn_font(run, CN_FONT)
            i = j
            continue

        # 表格：检测表头行 + 分隔行
        if stripped.startswith('|') and i + 1 < n and is_table_separator(lines[i+1]):
            # 收集表行直到非表行
            header = parse_table_row(lines[i])
            i += 2  # 跳过表头与分隔
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            # 填表头
            for c_i, c in enumerate(header):
                cell_p = table.rows[0].cells[c_i].paragraphs[0]
                parse_inline(cell_p, c)
            for r_i, row in enumerate(rows, start=1):
                for c_i, c in enumerate(row):
                    cell_p = table.rows[r_i].cells[c_i].paragraphs[0]
                    parse_inline(cell_p, c)
            style_table(table)
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', stripped):
            item = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, item)
            i += 1
            continue

        # 有序列表（1. 2.）
        if re.match(r'^\d+\.\s+', stripped):
            item = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, item)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print("saved:", docx_path)
    # 真实 Word 脚注注入
    nf = _inject_footnotes(docx_path)
    if nf:
        print("footnotes injected:", nf)
    # 后置批注转换钩子：将「批注（判断过程、思路与假设）」章节转为真实 Word 批注气泡
    _post_annotate(docx_path)


def _post_annotate(docx_path):
    """后置批注转换钩子（可选，由 ai-transparency-compliance skill 提供）。

    若文档含「批注（判断过程、思路与假设）」章节，则调用 annotations_to_docx_comments
    将其转为真实 Word 批注（comments.xml 气泡），并从正文移除该章节。
    转换器为幂等设计：文档无批注章节时返回 0、不改动文档。
    安全降级：转换器缺失/不可导入、或设置环境变量 MD2DOCX_NO_ANNOTATE 时静默跳过。
    """
    if os.environ.get('MD2DOCX_NO_ANNOTATE'):
        return
    try:
        import importlib.util
        candidates = [
            os.environ.get('ANNOTATE_SCRIPT'),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), 'scripts', 'annotations_to_docx_comments.py'),
            r'C:/Users/XY/.workbuddy/skills/ai-transparency-compliance/scripts/annotations_to_docx_comments.py',
        ]
        spec = None
        for c in candidates:
            if c and os.path.isfile(c):
                spec = importlib.util.spec_from_file_location('annotations_to_docx_comments', c)
                break
        if spec is None:
            return
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        n = mod.convert_docx_annotations(docx_path)
        if n:
            print("annotations -> Word comments:", n)
    except Exception as e:
        print("post_annotate skipped:", e)


if __name__ == '__main__':
    src = sys.argv[1]
    dst = sys.argv[2]
    convert(src, dst)
